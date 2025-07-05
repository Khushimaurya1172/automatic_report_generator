import tkinter as tk
from tkinter import messagebox
from docx import Document
from datetime import date
import os
import subprocess
import time

# Try importing docx2pdf
try:
    from docx2pdf import convert
    pdf_supported = True
except ImportError:
    pdf_supported = False

def generate_report():
    name = name_entry.get()
    role = role_entry.get()
    task = task_entry.get()
    title = title_entry.get()
    tech = tech_entry.get()
    outcome = outcome_text.get("1.0", "end").strip()

    if not (name and role and task and title and tech and outcome):
        messagebox.showerror("Error", "Please fill all fields.")
        return

    company = "Main Flow Services and Technologies Pvt. Ltd."
    today = date.today().strftime("%d-%m-%Y")

    doc = Document()
    doc.add_heading(f'{title} – Internship Report', 0)

    doc.add_paragraph(f"Date: {today}")
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Role: {role}")
    doc.add_paragraph(f"Company: {company}")
    doc.add_paragraph(f"Task Number: {task}")

    doc.add_heading("Project Title", level=1)
    doc.add_paragraph(title)

    doc.add_heading("Technologies Used", level=1)
    doc.add_paragraph(tech)

    doc.add_heading("Problem Statement", level=1)
    doc.add_paragraph("Generating intern reports manually takes time and leads to inconsistent formatting. "
                      "Every batch requires customized reports, which becomes a repetitive task for HR or coordinators.")

    doc.add_heading("Proposed Solution", level=1)
    doc.add_paragraph("This Python tool automates the generation of internship task reports in a structured Word document format. "
                      "It takes user input (name, task, date, etc.) and outputs a formatted .docx file ready for submission.")

    doc.add_heading("Outcomes & Benefits", level=1)
    doc.add_paragraph(outcome)

    doc.add_heading("Future Scope", level=1)
    doc.add_paragraph("This tool can be extended to include certificate generation, feedback collection forms, and PDF export support.")

    # Save path
    os.makedirs("reports", exist_ok=True)
    base_filename = f"Report_{name.replace(' ', '_')}"
    docx_path = os.path.join("reports", f"{base_filename}.docx")
    pdf_path = os.path.join("reports", f"{base_filename}.pdf")

    # Save DOCX
    if os.path.exists(docx_path):
        os.remove(docx_path)
    if os.path.exists(pdf_path):
        os.remove(pdf_path)

    doc.save(docx_path)

    # Convert to PDF
    if pdf_supported:
        try:
            convert(docx_path, pdf_path)
            messagebox.showinfo("Success", f"Report Saved:\n{pdf_path}")
            time.sleep(1)

            if os.path.exists(pdf_path):
                try:
                    subprocess.run(["code", pdf_path], shell=True)
                except Exception as e:
                    messagebox.showwarning("VS Code Open Failed", f"PDF created but VS Code open failed:\n{e}")
        except Exception as e:
            messagebox.showwarning("PDF Failed", f"DOCX saved but PDF conversion failed:\n{e}")
    else:
        messagebox.showinfo("Saved", f"DOCX saved at:\n{docx_path}\n(PDF not created)")

# GUI setup
root = tk.Tk()
root.title("Intern Report Generator")
root.geometry("600x500")

tk.Label(root, text="Name:").grid(row=0, column=0, sticky="w", padx=10, pady=5)
name_entry = tk.Entry(root, width=50)
name_entry.grid(row=0, column=1, pady=5)

tk.Label(root, text="Role:").grid(row=1, column=0, sticky="w", padx=10)
role_entry = tk.Entry(root, width=50)
role_entry.insert(0, "Python Developer Intern")
role_entry.grid(row=1, column=1)

tk.Label(root, text="Task Number:").grid(row=2, column=0, sticky="w", padx=10)
task_entry = tk.Entry(root, width=50)
task_entry.grid(row=2, column=1)

tk.Label(root, text="Project Title:").grid(row=3, column=0, sticky="w", padx=10)
title_entry = tk.Entry(root, width=50)
title_entry.grid(row=3, column=1)

tk.Label(root, text="Technologies Used:").grid(row=4, column=0, sticky="w", padx=10)
tech_entry = tk.Entry(root, width=50)
tech_entry.grid(row=4, column=1)

tk.Label(root, text="Outcome / Benefits:").grid(row=5, column=0, sticky="nw", padx=10)
outcome_text = tk.Text(root, height=5, width=38)
outcome_text.grid(row=5, column=1, pady=10)

generate_btn = tk.Button(root, text="Generate Report & PDF", command=generate_report, bg="#4CAF50", fg="white", width=25)
generate_btn.grid(row=6, column=1, pady=20)

root.mainloop()
